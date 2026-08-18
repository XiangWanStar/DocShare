# -*- coding: utf-8 -*-
"""SPA 全流程验证。"""
import io, json, time, httpx
from docx import Document
from playwright.sync_api import sync_playwright

APP = "http://127.0.0.1:8010"
client = httpx.Client(base_url=APP, timeout=30)
ok = fail = 0
def check(name, cond, extra=""):
    global ok, fail
    if cond: ok += 1; print("  [PASS]", name)
    else: fail += 1; print("  [FAIL]", name, extra)

email = "spa" + str(int(time.time())) + "@example.com"

with sync_playwright() as p:
    browser = p.chromium.launch(channel="msedge", headless=True, args=["--incognito", "--disable-extensions"])
    ctx = browser.new_context(viewport={"width": 1440, "height": 900})
    page = ctx.new_page()
    errors = []
    page.on("pageerror", lambda e: errors.append(str(e)[:200]))
    page.on("console", lambda m: errors.append(f"[{m.type}] {m.text[:150]}") if m.type == "error" else None)

    print("== 1. 根路径重定向到 SPA ==")
    r = client.get("/", follow_redirects=False)
    check("GET / → 307", r.status_code == 307, str(r.status_code))
    check("重定向到 /static/vue/", "/static/vue/" in r.headers.get("location", ""), r.headers.get("location"))

    print("== 2. 未登录访问 → 跳转登录 ==")
    page.goto(APP + "/static/vue/", wait_until="domcontentloaded")
    page.wait_for_timeout(2500)
    check("SPA 加载并重定向到 #/login", page.url.endswith("#/login"), page.url)
    check("登录页可见", page.locator("text=登 录").count() > 0)

    print("== 3. 注册 ==")
    page.goto(APP + "/static/vue/#/register", wait_until="domcontentloaded")
    page.wait_for_timeout(1000)
    page.fill("input[placeholder='昵称（可选）']", "SPA用户")
    page.fill("input[placeholder='请输入邮箱']", email)
    page.fill("input[placeholder='设置密码（至少6位）']", "secret123")
    page.click("button:has-text('注 册')")
    page.wait_for_url("**/#/dashboard", timeout=15000)
    check("注册后进入 dashboard", "/dashboard" in page.url, page.url)

    print("== 4. 文件列表 + 上传 ==")
    page.wait_for_timeout(3000)
    empty_visible = page.locator("text=暂无文件，点击右上角上传").count() > 0
    check("空状态提示可见", empty_visible)
    # 生成真实 docx 并上传（新交互：上传模态框）
    doc = Document()
    doc.add_heading("SPA 测试文档", level=1)
    doc.add_paragraph("通过 Vue 前端上传的文档。")
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    page.click(".btn-upload")
    page.wait_for_timeout(500)
    page.set_input_files(".upload-modal input[type=file]", {"name": "SPA测试.docx", "mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "buffer": buf.getvalue()})
    page.wait_for_timeout(600)
    page.click("button:has-text('开始上传')")
    ok_tip = False
    for _ in range(12):
        page.wait_for_timeout(1000)
        if page.locator("text=成功上传 1 个文件").count() > 0: ok_tip = True; break
    check("成功提示", ok_tip)
    page.wait_for_timeout(2000)
    check("上传后列表出现文件", page.locator("text=SPA测试.docx").count() > 0)

    print("== 5. 打开查看器 ==")
    page.click("text=SPA测试.docx", timeout=10000)
    page.wait_for_timeout(3000)
    check("进入查看器路由", "/viewer" in page.url, page.url)
    # 等 OnlyOffice iframe + 高度修复
    ifr_ok = False
    for _ in range(12):
        page.wait_for_timeout(5000)
        if page.locator('iframe[src*="web-apps"]').count() > 0:
            box = page.locator('iframe[src*="web-apps"]').bounding_box()
            if box and box["height"] > 500:
                ifr_ok = True
                check("编辑器 iframe 高度撑满", box["height"] > 700, str(box))
                break
    if not ifr_ok:
        check("编辑器 iframe 高度撑满", False)

    print("== 6. 创建分享 ==")
    page.goto(APP + "/static/vue/#/dashboard", wait_until="domcontentloaded")
    page.wait_for_timeout(3000)
    # 网格视图卡片上的分享按钮
    share_btn = page.locator(".file-card .card-act[title='分享']").first
    if share_btn.count() == 0:
        # 切到列表视图
        page.click("button[title='列表视图']")
        page.wait_for_timeout(500)
        page.locator("button[title='分享']").first.click()
    else:
        share_btn.click()
    page.wait_for_timeout(800)
    check("分享模态框打开", page.locator("text=创建分享链接").count() > 0)
    page.click("button:has-text('创 建')") if page.locator("button:has-text('创 建')").count() else page.click(".modal-footer .btn-primary")
    page.wait_for_timeout(2000)
    share_url_el = page.locator("text=/http.*\/s\//").first
    check("分享链接生成", share_url_el.count() > 0, "")

    # 提取分享 URL
    share_url = ""
    if share_url_el.count():
        share_url = share_url_el.inner_text().strip()
    print("share_url:", share_url)

    print("== 7. 分享落地页 ==")
    if share_url:
        token = share_url.rstrip("/").split("/s/")[-1]
        r2 = client.get("/s/" + token, follow_redirects=False)
        check("GET /s/{token} → 307 到 SPA", r2.status_code == 307 and "#/share/" in r2.headers.get("location", ""), str(r2.status_code) + r2.headers.get("location", ""))
        page.goto(APP + "/static/vue/#/share/" + token, wait_until="domcontentloaded")
        # 无密码 → 自动进入编辑器
        ifr_ok2 = False
        for _ in range(10):
            page.wait_for_timeout(3000)
            if page.locator('iframe[src*="web-apps"]').count() > 0:
                ifr_ok2 = True
                break
        check("分享页编辑器 iframe 加载", ifr_ok2)

    print("== 8. 页面 JS 错误 ==")
    real_errors = [e for e in errors if "404" not in e]
    check("无 JS 错误", len(real_errors) == 0, str(real_errors[:3]))
    browser.close()

print(f"\n===== SPA 全流程: {ok} 通过, {fail} 失败 =====")
import sys; sys.exit(1 if fail else 0)